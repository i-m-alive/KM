"""DB-backed integration tests (item 3a/3c) - the counterpart to
regression_check.py for logic that genuinely needs a real database session
rather than a fake/mocked one. regression_check.py deliberately stays
DB-free (fast, safe to run anywhere); the tests below exercise the exact
ORM query paths that DB-free suite cannot: dictionary.find_in_text's
substring sweep against real MaskingEntity/MaskingAlias rows (Task A's alt-
text fix), dictionary.validate_custom_replacement's real
`db.query(MaskingEntity).filter(...)` (Task C), and the cross-surface
entity-merge flow (item 1a) via dictionary.add_alias.

This does NOT run the full detect()/apply() pipeline end to end - doing
that would require mocking the LLM detector, image scan, and summarizer
Bedrock calls in addition to the DB, which is a much bigger undertaking than
what's actually untested here. What's tested is the real database layer
underneath the fixes that were only ever exercised through fakes before.

SAFETY: connects to a SEPARATE, disposable database (naviknow_test by
default, override with NAVIKNOW_TEST_DATABASE_URL) - never the real
configured app database. Creates its own rows and cleans up after itself
via a transaction ROLLBACK, never a commit, so nothing written here persists
even in the test database.

Setup (once): create the test database and load the schema -
    createdb -U naviknow -h localhost -p 5433 naviknow_test
    psql -U naviknow -h localhost -p 5433 -d naviknow_test -f backend/init.sql

Run:
    cd backend && .venv/bin/python scripts/db_integration_check.py
"""

import os
import sys
import uuid
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

TEST_DATABASE_URL = os.environ.get(
    "NAVIKNOW_TEST_DATABASE_URL", "postgresql+psycopg://naviknow:naviknow@localhost:5433/naviknow_test"
)

failures: list[str] = []
passes = 0


def check(name: str, ok: bool, detail: str = "") -> None:
    global passes
    if ok:
        passes += 1
        print(f"  ok    {name}")
    else:
        failures.append(f"{name}: {detail}")
        print(f"  FAIL  {name}  {detail}")


def _make_session():
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker

    engine = create_engine(TEST_DATABASE_URL)
    return sessionmaker(bind=engine)()


def _seed_run(db):
    """MaskingEntity.created_by_run_id is a real FK to agent_runs, which in
    turn requires a real user - dictionary.get_or_create can't be exercised
    against a real DB without a real (if throwaway) run to attribute it to.
    Never committed - the caller's rollback discards this along with
    everything else."""
    from app.models import AgentRun, Role, User

    role = db.query(Role).filter(Role.name == "admin").first()
    user = User(email=f"db-integration-{uuid.uuid4()}@test.local", hashed_password="x", role_id=role.id)
    db.add(user)
    db.flush()
    run = AgentRun(agent_id="sanitization", status="pending", input_json={}, created_by=user.id)
    db.add(run)
    db.flush()
    return run


def check_find_in_text_real_db() -> None:
    """Task A's actual fix: dictionary.find_in_text must find a known,
    APPROVED entity's alias as a SUBSTRING of a longer string (e.g.
    "NextCare" inside "Nextcare_logo" - the real production leak), using a
    real query against real MaskingEntity/MaskingAlias rows, not a fake."""
    print("\n== dictionary.find_in_text against a real DB (item 3a) ==")
    from app.masking import dictionary

    db = _make_session()
    try:
        run_id = _seed_run(db).id
        entity = dictionary.get_or_create(db, "NextCare", "CLIENT_NAME", run_id, approved=True)
        dictionary.approve(db, entity)

        hits = dictionary.find_in_text(db, "See the descr Nextcare_logo.png attached")
        check("find_in_text matches a known entity embedded in a longer string (case-insensitive)",
              any(e.id == entity.id for e, _ in hits), f"got: {[(e.mask_token, s) for e, s in hits]}")

        no_hits = dictionary.find_in_text(db, "Nothing relevant here at all")
        check("find_in_text returns nothing when the text doesn't contain any known entity",
              not any(e.id == entity.id for e, _ in no_hits))
    finally:
        db.rollback()
        db.close()


def check_alias_validation_real_db() -> None:
    """Task C's validate_custom_replacement against REAL MaskingEntity rows
    and a real query - the offline suite only ever exercised this through a
    hand-rolled fake query object."""
    print("\n== dictionary.validate_custom_replacement against a real DB (item 3a) ==")
    from app.masking import dictionary

    db = _make_session()
    try:
        run_id = _seed_run(db).id
        target = dictionary.get_or_create(db, "Johnson & Johnson", "CLIENT_NAME", run_id, approved=True)
        dictionary.approve(db, target)
        other = dictionary.get_or_create(db, "Pfizer", "CLIENT_NAME", run_id, approved=True)
        dictionary.approve(db, other)

        problems = dictionary.validate_custom_replacement(db, target, "Pfizer")
        check("aliasing one real entity to another tracked entity's own name is rejected",
              len(problems) > 0, f"got: {problems}")

        clean = dictionary.validate_custom_replacement(db, target, "Acme Pharma")
        check("a genuinely distinct, untracked alias passes", clean == [], f"got: {clean}")

        dictionary.set_custom_replacement(db, target, "Acme Pharma")
        db.flush()
        dup_problems = dictionary.validate_custom_replacement(db, other, "Acme Pharma")
        check("an alias already assigned to a different real entity is rejected",
              len(dup_problems) > 0, f"got: {dup_problems}")
    finally:
        db.rollback()
        db.close()


def check_cross_surface_merge_real_db() -> None:
    """Item 1a's cross-surface merge (edits.entity_merges): merging one
    surface into another's entity via dictionary.add_alias must make BOTH
    surfaces resolve to the SAME token/alias afterward, using real rows."""
    print("\n== cross-surface entity merge against a real DB (item 3a) ==")
    from app.masking import dictionary

    db = _make_session()
    try:
        run_id = _seed_run(db).id
        canonical = dictionary.get_or_create(db, "Johnson & Johnson India", "CLIENT_NAME", run_id, approved=True)
        dictionary.approve(db, canonical)

        # Simulate agent.py's merge branch: "J&J" merges into the canonical
        # entity instead of getting its own get_or_create call.
        dictionary.add_alias(db, canonical, "J&J")
        db.flush()

        looked_up = dictionary.lookup(db, "J&J")
        check("the merged surface now resolves to the canonical entity via lookup()",
              looked_up is not None and looked_up.id == canonical.id, f"got: {looked_up}")
        check("resolved_replacement is identical for both surfaces (same token)",
              dictionary.resolved_replacement(canonical) == dictionary.resolved_replacement(looked_up))
    finally:
        db.rollback()
        db.close()


def check_reapply_with_additional_entities_real_db() -> None:
    """revalidate.py's fix loop against a real DB + real file I/O - the
    riskiest new code path (writes new MaskingEntity/MaskingOccurrence rows
    AND mutates a file on disk). Builds a minimal docx that already has one
    entity masked (simulating a completed run's rendered output) plus one
    UNMASKED residual name - exactly the "Arvind Fashions" production miss
    this whole feature exists to catch - then approves that residual via
    reapply_with_additional_entities and checks it actually disappears from
    the file, gets a real dictionary entry, and leaves the pre-existing
    mask token untouched (idempotent re-render, same as remediate.py's)."""
    print("\n== revalidate.reapply_with_additional_entities against a real DB + real file ==")
    import shutil
    import tempfile

    import docx

    from app.agents.sanitization import revalidate
    from app.documents.extract import extract_chunks
    from app.masking import dictionary
    from app.models import AgentRun, MaskingAlias, MaskingEntity, MaskingOccurrence, UploadedDocument, User

    db = _make_session()
    workdir = None
    try:
        run = _seed_run(db)
        run_id = run.id
        user_id = run.created_by

        known_entity = dictionary.get_or_create(db, "Johnson & Johnson", "CLIENT_NAME", run_id, approved=True)
        dictionary.approve(db, known_entity)
        db.flush()
        db.add(MaskingOccurrence(
            run_id=run_id, entity_id=known_entity.id, chunk_id=0, start_offset=0, end_offset=0,
            surface_text="Johnson & Johnson",
        ))

        workdir = tempfile.mkdtemp(prefix="naviknow-db-revalidate-")
        masked_path = f"{workdir}/masked.docx"
        d = docx.Document()
        d.add_paragraph(f"Engagement overview for {known_entity.mask_token}.")
        d.add_paragraph("D2C website partnership announced with Arvind Fashions this quarter.")
        d.save(masked_path)

        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        doc = UploadedDocument(
            filename="engagement.docx", content_type=content_type,
            stored_path=masked_path, uploaded_by=run.created_by,
        )
        db.add(doc)
        db.flush()

        run.output_json = {
            "document_id": str(doc.id),
            "masked_document_path": masked_path,
            "masking_style": "token",
            "verified_images": True,
        }
        db.flush()

        result = revalidate.reapply_with_additional_entities(
            db, run, [{"leaked_text": "Arvind Fashions", "entity_type": "CLIENT_NAME"}]
        )
        check("reapply reports exactly one new entity applied", result["applied"] == 1, f"got: {result}")

        new_entity = dictionary.lookup(db, "Arvind Fashions")
        check("a real, approved MaskingEntity now exists for the approved residual",
              new_entity is not None and new_entity.status == "approved", f"got: {new_entity}")

        chunks = extract_chunks(masked_path, content_type, "engagement.docx")
        full_text = "\n".join(c.text for c in chunks)
        check("the approved residual no longer appears in the rendered file",
              "Arvind Fashions" not in full_text, f"got: {full_text!r}")
        check("the PRE-EXISTING masked entity's token is untouched (idempotent re-render)",
              known_entity.mask_token in full_text, f"got: {full_text!r}")
        check("run.output_json now carries a revalidation block for the fix pass",
              "revalidation" in (run.output_json or {}), f"got keys: {list((run.output_json or {}).keys())}")
    finally:
        # reapply_with_additional_entities COMMITS internally - it MUST, in
        # production, for the exact same reason remediate_run() does (real
        # work has to persist) - so unlike every other check in this file
        # (whose functions only flush), a plain db.rollback() here has
        # nothing left to undo. Explicit cleanup, in FK-safe order, so this
        # disposable test DB doesn't accumulate rows run after run.
        try:
            db.rollback()
            new_entity = dictionary.lookup(db, "Arvind Fashions")
            entity_ids = [e.id for e in (locals().get("known_entity"), new_entity) if e is not None]
            if entity_ids:
                db.query(MaskingOccurrence).filter(MaskingOccurrence.entity_id.in_(entity_ids)).delete(synchronize_session=False)
                db.query(MaskingAlias).filter(MaskingAlias.entity_id.in_(entity_ids)).delete(synchronize_session=False)
                db.query(MaskingEntity).filter(MaskingEntity.id.in_(entity_ids)).delete(synchronize_session=False)
            db.query(MaskingOccurrence).filter(MaskingOccurrence.run_id == run_id).delete(synchronize_session=False)
            db.query(UploadedDocument).filter(UploadedDocument.id == doc.id).delete(synchronize_session=False)
            db.query(AgentRun).filter(AgentRun.id == run_id).delete(synchronize_session=False)
            db.query(User).filter(User.id == user_id).delete(synchronize_session=False)
            db.commit()
        except Exception:
            db.rollback()
        db.close()
        if workdir:
            shutil.rmtree(workdir, ignore_errors=True)


def main() -> int:
    try:
        from sqlalchemy import create_engine

        create_engine(TEST_DATABASE_URL).connect().close()
    except Exception as exc:
        print(f"Could not connect to the test database at {TEST_DATABASE_URL}: {exc}")
        print("Setup: createdb naviknow_test, then run backend/init.sql against it (see this file's docstring).")
        return 2

    for fn in (
        check_find_in_text_real_db, check_alias_validation_real_db, check_cross_surface_merge_real_db,
        check_reapply_with_additional_entities_real_db,
    ):
        try:
            fn()
        except Exception as exc:
            import traceback

            traceback.print_exc()
            failures.append(f"{fn.__name__}: crashed - {exc}")

    print(f"\n{'=' * 50}\n{passes} checks passed, {len(failures)} failed")
    for f in failures:
        print(f"  FAIL  {f}")
    return 1 if failures else 0


if __name__ == "__main__":
    sys.exit(main())
